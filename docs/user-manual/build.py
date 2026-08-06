"""
Build HR_Intelligence_Platform_User_Manual.docx + matching .pdf.

The PDF is exported from the Word file (Microsoft Word COM) so both formats
always share the same content, structure, and screenshots.

Usage:
  python docs/user-manual/capture.py
  python docs/user-manual/build.py
"""
from __future__ import annotations

import json
from collections import OrderedDict
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
MANIFEST_PATH = SHOTS / "manifest.json"
VERSION = "1.0"
TODAY = date.today().isoformat()
COMPANY = "Techberry Infotech Pvt. Ltd."
DOC_NAME = "HR_Intelligence_Platform_User_Manual"

# Chapter meta for modules (only implemented surfaces)
MODULE_META = OrderedDict(
    [
        (
            "01-home",
            {
                "title": "Home",
                "purpose": "Public landing experience that introduces the product and routes users to Jobs.",
                "nav": "/",
                "roles": "Public",
            },
        ),
        (
            "02-authentication",
            {
                "title": "Authentication",
                "purpose": (
                    "Staff login and password recovery for Recruiter, Head HR, and CEO accounts. "
                    "Forgot password sends a 6-digit OTP (valid 10 minutes), then verify OTP and set a new password. "
                    "New staff accounts are created by Head HR (Admins), not via public self-signup."
                ),
                "nav": "/login → /login/admin → /forgot-password/admin → verify OTP → reset password",
                "roles": "Public (unauthenticated); after login: RECRUITER | HEAD_HR | CEO",
            },
        ),
        (
            "03-public-jobs",
            {
                "title": "Jobs (Public Board & Apply)",
                "purpose": "Browse published jobs and submit applications via the apply modal.",
                "nav": "/jobs",
                "roles": "Public (apply); staff may browse but cannot apply while signed in as staff",
            },
        ),
        (
            "04-support",
            {
                "title": "Support",
                "purpose": "FAQ, Contact Us, and HRMS Testing Feedback channels.",
                "nav": "/support/faq · /support/contact · /support/hrms-feedback",
                "roles": "Public",
            },
        ),
        (
            "05-head-hr-overview",
            {
                "title": "Head HR Overview Dashboard",
                "purpose": "Organization overview and Head HR panel entry point.",
                "nav": "/head-hr",
                "roles": "HEAD_HR",
            },
        ),
        (
            "06-head-hr-admins",
            {
                "title": "Admins",
                "purpose": "Create and manage recruiter/admin HR accounts for the organization.",
                "nav": "/head-hr/admins",
                "roles": "HEAD_HR",
            },
        ),
        (
            "07-head-hr-candidates",
            {
                "title": "Head HR Candidates",
                "purpose": "Org-wide candidate list and candidate detail.",
                "nav": "/head-hr/candidates · /head-hr/candidates/:cid",
                "roles": "HEAD_HR (CEO uses /ceo/candidates read-only)",
            },
        ),
        (
            "08-head-hr-jobs",
            {
                "title": "Head HR Jobs",
                "purpose": "Organization job inventory with search and management actions.",
                "nav": "/head-hr/jobs",
                "roles": "HEAD_HR",
            },
        ),
        (
            "09-head-hr-job-detail",
            {
                "title": "Head HR Job Details & Applied Candidates",
                "purpose": "Inspect a job and its applicants.",
                "nav": "/head-hr/jobs/:jdid",
                "roles": "HEAD_HR (CEO: /ceo/jobs/:jdid read-only)",
            },
        ),
        (
            "10-candidate-evaluation",
            {
                "title": "Candidate Evaluation (Profile & Match)",
                "purpose": "Review applicant profile and ATS match analysis for a job application.",
                "nav": "/head-hr/jobs/:jdid/candidates/:cid",
                "roles": "HEAD_HR; CEO read-only under /ceo/…",
            },
        ),
        (
            "11-bulk-parsing",
            {
                "title": "Bulk Resume Parser (Head HR)",
                "purpose": "Batch-parse resumes from ZIP/files.",
                "nav": "/head-hr/bulk-parsing",
                "roles": "HEAD_HR",
            },
        ),
        (
            "12-integrations",
            {
                "title": "Integrations (Head HR)",
                "purpose": "Monitor external publishing integrations.",
                "nav": "/head-hr/integrations",
                "roles": "HEAD_HR",
            },
        ),
        (
            "13-settings",
            {
                "title": "Head HR Settings",
                "purpose": "Security (password) and integrations settings for Head HR.",
                "nav": "/head-hr/settings",
                "roles": "HEAD_HR",
            },
        ),
        (
            "14-recruiter-dashboard",
            {
                "title": "Recruiter Dashboard",
                "purpose": "Recruiter home for creating, editing, publishing, and managing own jobs.",
                "nav": "/dashboard",
                "roles": "RECRUITER",
            },
        ),
        (
            "15-recruiter-candidates",
            {
                "title": "Applied Candidates (Recruiter)",
                "purpose": "Shortlist/reject applicants and view match reasons for recruiter jobs.",
                "nav": "/candidates",
                "roles": "RECRUITER",
            },
        ),
        (
            "16-recruiter-bulk",
            {
                "title": "Bulk Resume Parser (Recruiter)",
                "purpose": "Recruiter entry to the bulk resume parser.",
                "nav": "/admin/bulk-resume-parser",
                "roles": "RECRUITER",
            },
        ),
        (
            "17-feedback-admin",
            {
                "title": "Feedback Admin",
                "purpose": "Review HRMS testing feedback submissions.",
                "nav": "/admin/feedback",
                "roles": "RECRUITER",
            },
        ),
        (
            "18-recruiter-integrations",
            {
                "title": "Integrations (Staff)",
                "purpose": "Integrations dashboard available to signed-in staff via navbar/settings routes.",
                "nav": "/integrations",
                "roles": "RECRUITER | HEAD_HR | CEO",
            },
        ),
        (
            "19-recruiter-settings",
            {
                "title": "Settings (Staff)",
                "purpose": "Account security settings for signed-in staff.",
                "nav": "/settings",
                "roles": "RECRUITER | HEAD_HR | CEO",
            },
        ),
        (
            "20-ceo-overview",
            {
                "title": "CEO / Executive Dashboard",
                "purpose": "Read-only executive views of overview, jobs, candidates, and evaluation.",
                "nav": "/ceo · /ceo/jobs · /ceo/candidates",
                "roles": "CEO",
            },
        ),
        (
            "21-logout",
            {
                "title": "Logout",
                "purpose": "End the authenticated session and return to login.",
                "nav": "Sidebar Logout (Head HR/CEO) or Navbar Logout (Recruiter)",
                "roles": "RECRUITER | HEAD_HR | CEO",
            },
        ),
    ]
)


def load_manifest() -> list[dict]:
    if not MANIFEST_PATH.exists():
        raise SystemExit("Run capture.py first (screenshots/manifest.json missing).")
    return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))


def by_module(manifest: list[dict]) -> OrderedDict[str, list[dict]]:
    out: OrderedDict[str, list[dict]] = OrderedDict()
    for key in MODULE_META:
        out[key] = []
    for item in manifest:
        out.setdefault(item["module"], []).append(item)
    return out


def _set_run_font(run, size=11, bold=False, color=None):
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


_BOOKMARK_SEQ = 0


def _next_bookmark_id() -> int:
    global _BOOKMARK_SEQ
    _BOOKMARK_SEQ += 1
    return _BOOKMARK_SEQ


def _add_bookmark(paragraph, name: str) -> None:
    """Attach a Word bookmark to a paragraph (for TOC jump targets)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bid = str(_next_bookmark_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph, anchor: str, text: str, *, size_pt: int = 10) -> None:
    """Add a clickable internal hyperlink (Word + PDF after Word export)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0369A1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size_pt * 2))
    r_pr.append(sz_cs)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_pageref_field(paragraph, bookmark: str, *, size_pt: int = 10) -> None:
    """Insert a PAGEREF field (hyperlinked) so TOC page numbers update in Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt as DocPt
    from docx.shared import RGBColor

    run = paragraph.add_run()
    run.font.name = "Calibri"
    run.font.size = DocPt(size_pt)
    run.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "…"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(placeholder)
    run._r.append(end)


def build_toc_entries(manifest: list[dict]) -> list[tuple[str, str, str]]:
    """Return [(chapter_no, title, bookmark), ...] for TOC table."""
    entries: list[tuple[str, str, str]] = [("1", "Introduction", "sec_1")]
    modules = by_module(manifest)
    chapter = 2
    for mod_key, meta in MODULE_META.items():
        if modules.get(mod_key):
            entries.append((str(chapter), meta["title"], f"sec_{chapter}"))
            chapter += 1
    entries.extend(
        [
            (str(chapter), "Forms, Fields, and Buttons Reference", f"sec_{chapter}"),
            (str(chapter + 1), "Troubleshooting", f"sec_{chapter + 1}"),
            (str(chapter + 2), "FAQs", f"sec_{chapter + 2}"),
            (str(chapter + 3), "Appendix", f"sec_{chapter + 3}"),
        ]
    )
    return entries


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    from docx.shared import Pt as DocPt

    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = DocPt(3)
    para.paragraph_format.space_after = DocPt(3)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = DocPt(size)


def _set_cell_toc_link(cell, text: str, bookmark: str, *, size: int = 10) -> None:
    """TOC section cell: clickable jump to chapter bookmark."""
    from docx.shared import Pt as DocPt

    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = DocPt(3)
    para.paragraph_format.space_after = DocPt(3)
    _add_internal_hyperlink(para, bookmark, text, size_pt=size)


def _set_cell_pageref(cell, bookmark: str, *, size: int = 10) -> None:
    """TOC page cell: PAGEREF field (also clickable)."""
    from docx.shared import Pt as DocPt

    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = DocPt(3)
    para.paragraph_format.space_after = DocPt(3)
    _add_pageref_field(para, bookmark, size_pt=size)


def build_docx(manifest: list[dict]) -> Path:
    global _BOOKMARK_SEQ
    _BOOKMARK_SEQ = 0
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size, space_before, space_after in (
        (1, 24, 0, 14),
        (2, 16, 16, 10),
        (3, 13, 12, 8),
    ):
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Calibri"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            hs.paragraph_format.space_before = Pt(space_before)
            hs.paragraph_format.space_after = Pt(space_after)
            hs.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass

    def configure_section(section, *, header: bool):
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        if header:
            hp = section.header.paragraphs[0]
            hp.text = ""
            r = hp.add_run("HR Intelligence Platform  |  User Manual")
            _set_run_font(r, size=9, color=(0x64, 0x74, 0x8B))
            fp = section.footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.text = ""
            r = fp.add_run(f"Confidential  ·  {COMPANY}  ·  v{VERSION}  ·  Page ")
            _set_run_font(r, size=9, color=(0x64, 0x74, 0x8B))
            _add_page_number(fp)
        else:
            section.header.paragraphs[0].text = ""
            section.footer.paragraphs[0].text = ""

    def h(text, level=1, bookmark=None):
        para = doc.add_heading(text, level=level)
        # Enforce size + gap even if Word theme overrides the style
        sizes = {1: 24, 2: 16, 3: 13}
        gaps_after = {1: 14, 2: 10, 3: 8}
        gaps_before = {1: 0, 2: 16, 3: 12}
        para.paragraph_format.space_before = Pt(gaps_before.get(level, 0))
        para.paragraph_format.space_after = Pt(gaps_after.get(level, 8))
        for run in para.runs:
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(sizes.get(level, 14))
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        if bookmark:
            _add_bookmark(para, bookmark)
        return para

    def p(text, bold=False, center=False, *, keep_with_next=False, page_break_before=False):
        para = doc.add_paragraph()
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if page_break_before:
            para.paragraph_format.page_break_before = True
        if keep_with_next:
            para.paragraph_format.keep_with_next = True
        if bold:
            # Inline section labels (e.g. "How to use this manual")
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            r = para.add_run(text)
            _set_run_font(r, size=13, bold=True, color=(0x0F, 0x17, 0x2A))
        else:
            para.paragraph_format.space_after = Pt(10)
            r = para.add_run(text)
            _set_run_font(r, size=11, bold=False)
        return para

    def bullet(text, *, keep_with_next=False):
        para = doc.add_paragraph(text, style="List Bullet")
        if keep_with_next:
            para.paragraph_format.keep_with_next = True
        return para

    def add_figure(item: dict, fig_id: str):
        path = SHOTS / item["file"]
        if not path.exists():
            p(f"[Missing figure: {item['file']}]")
            return
        doc.add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure {fig_id}  {item['title']}")
        r.italic = True
        _set_run_font(r, size=9, color=(0x47, 0x55, 0x69))
        return cap

    # Cover
    configure_section(doc.sections[0], header=False)
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HR Intelligence Platform")
    _set_run_font(r, size=32, bold=True, color=(0x0F, 0x17, 0x2A))
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("User Manual")
    _set_run_font(r, size=24, color=(0x03, 0x69, 0xA1))
    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Version {VERSION}\n{TODAY}\n{COMPANY}\nEnterprise Customer Documentation\n"
        f"Document ID: {DOC_NAME}"
    )
    _set_run_font(r, size=12, color=(0x64, 0x74, 0x8B))
    for _ in range(5):
        doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("CONFIDENTIAL — Documents only features implemented in this release")
    _set_run_font(r, size=9, color=(0x94, 0xA3, 0xB8))

    # History + TOC section
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, header=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    r = title.add_run("Document History")
    _set_run_font(r, size=24, bold=True, color=(0x0F, 0x17, 0x2A))
    p("Revision history for this customer-facing user manual.")
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    for i, txt in enumerate(("Version", "Author", "Date", "Description")):
        _set_cell_text(table.rows[0].cells[i], txt, bold=True, size=10)
        _shade_cell(table.rows[0].cells[i], "E2E8F0")
    _set_cell_text(table.rows[1].cells[0], VERSION, size=10)
    _set_cell_text(table.rows[1].cells[1], "Technical Publications", size=10)
    _set_cell_text(table.rows[1].cells[2], TODAY, size=10)
    _set_cell_text(
        table.rows[1].cells[3],
        "Complete enterprise manual rebuilt from live routes and screenshots",
        size=10,
    )
    doc.add_page_break()

    # ——— Table of Contents (formatted table) ———
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_after = Pt(14)
    r = toc_title.add_run("Table of Contents")
    _set_run_font(r, size=24, bold=True, color=(0x0F, 0x17, 0x2A))
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    r = intro.add_run(
        "Click any section name (or page number) to jump to that chapter."
    )
    _set_run_font(r, size=10, color=(0x64, 0x74, 0x8B))

    toc_entries = build_toc_entries(manifest)
    toc_table = doc.add_table(rows=1 + len(toc_entries), cols=3)
    toc_table.style = "Table Grid"
    for i, hdr in enumerate(("No.", "Section", "Page")):
        _set_cell_text(toc_table.rows[0].cells[i], hdr, bold=True, size=10)
        _shade_cell(toc_table.rows[0].cells[i], "0369A1")
        for run in toc_table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_i, (num, section, bookmark) in enumerate(toc_entries, start=1):
        _set_cell_text(toc_table.rows[row_i].cells[0], num, size=10)
        _set_cell_toc_link(toc_table.rows[row_i].cells[1], section, bookmark, size=10)
        _set_cell_pageref(toc_table.rows[row_i].cells[2], bookmark, size=10)
        if row_i % 2 == 0:
            for c in range(3):
                _shade_cell(toc_table.rows[row_i].cells[c], "F8FAFC")
    # Set column widths
    for row in toc_table.rows:
        row.cells[0].width = Inches(0.7)
        row.cells[1].width = Inches(4.8)
        row.cells[2].width = Inches(0.8)
    doc.add_page_break()

    # ——— Introduction (single page) → then Home modules ———
    h("1. Introduction", 1, bookmark="sec_1")
    p(
        "HR Intelligence Platform is an AI-assisted recruitment application for publishing jobs, "
        "receiving applications, parsing resumes and job descriptions, scoring candidate–job fit, "
        "and enabling Recruiters, Head of HR, and Executives (CEO) to evaluate applicants."
    )
    p(
        "This user manual walks through every shipped screen with live screenshots. "
        "Follow each chapter in order, or jump from the Table of Contents to a specific module. "
        "Only features that exist in the current product are documented."
    )
    p("How to use this manual", bold=True)
    bullet("Start with Home, then Authentication, then the modules for your role (Recruiter, Head HR, or CEO).")
    bullet("Each chapter lists Purpose, Navigation Path, Accessible Roles, then step-by-step actions with screenshots.")
    bullet("Supported staff roles: RECRUITER, HEAD_HR, and CEO (read-only). Admin Login authenticates these roles.")
    bullet("Recommended browsers: Chrome, Edge, Firefox, or Safari (latest). Desktop resolution 1280×720 or higher.")
    doc.add_page_break()

    modules = by_module(manifest)
    chapter = 2
    for mod_key, meta in MODULE_META.items():
        items = modules.get(mod_key) or []
        if not items:
            continue
        h(f"{chapter}. {meta['title']}", 1, bookmark=f"sec_{chapter}")
        p("Purpose", bold=True)
        p(meta["purpose"])
        p("Navigation Path", bold=True)
        p(meta["nav"])
        p("Accessible Roles", bold=True)
        p(meta["roles"])
        p("Description / Business Purpose", bold=True)
        p(
            f"This chapter documents the live UI for {meta['title']}. "
            "Each step includes a full screenshot captured from the running application."
        )

        for i, item in enumerate(items, start=1):
            h(f"Step {i}", 3)
            p("Action:", bold=True, keep_with_next=True)
            bullet(item["action"])
            add_figure(item, f"{chapter}.{i}")
            # Keep Description + Expected Result together on the page after the screenshot
            p("Description:", bold=True, page_break_before=True, keep_with_next=True)
            bullet(item.get("title") or item["action"], keep_with_next=True)
            p("Expected Result:", bold=True, keep_with_next=True)
            bullet(item["expected"])

        p("Notes", bold=True)
        bullet("If a control is disabled, complete required fields first.")
        bullet("Match scores depend on parsed JD and resume content.")
        doc.add_page_break()
        chapter += 1

    # Reference chapters — continue from next chapter number
    # (placeholder removed; existing block below still uses chapter variable)
    h(f"{chapter}. Forms, Fields, and Buttons Reference", 1, bookmark=f"sec_{chapter}")
    p("The following reference covers controls present in the shipped UI.")
    p("Job create / edit (Recruiter dashboard & Head HR overview posting)", bold=True)
    fields = [
        ("Job Title", "Required display name of the role"),
        ("Company", "Usually from the staff account company"),
        ("Location", "City / region / remote indicator"),
        ("Salary", "Optional compensation text"),
        ("Experience Range", "Minimum and maximum years"),
        ("Required / Mandatory Skills", "Used by ATS skills gate"),
        ("Preferred Skills", "Nice-to-have skills"),
        ("Description", "Full job description"),
        ("Keywords", "Search/match context keywords"),
    ]
    ft = doc.add_table(rows=1 + len(fields), cols=2)
    ft.style = "Table Grid"
    ft.rows[0].cells[0].text = "Field"
    ft.rows[0].cells[1].text = "Purpose / notes"
    for i, (a, b) in enumerate(fields, start=1):
        ft.rows[i].cells[0].text = a
        ft.rows[i].cells[1].text = b
    p("")
    p("Primary buttons", bold=True)
    actions = [
        ("Sign in", "Authenticate staff at /login/admin"),
        ("Get Started", "Navigate from landing to /jobs"),
        ("Apply", "Open apply modal on public jobs board"),
        ("Preview Post / Post Job", "Preview and publish a job (Recruiter / Head HR)"),
        ("Shortlist / Reject", "Application status actions (Recruiter Applied Candidates)"),
        ("Create Admin", "Provision HR account (Head HR Admins)"),
        ("Upload / Parse", "Bulk resume parser actions"),
        ("Logout", "End session"),
    ]
    at = doc.add_table(rows=1 + len(actions), cols=2)
    at.style = "Table Grid"
    at.rows[0].cells[0].text = "Button"
    at.rows[0].cells[1].text = "Behavior"
    for i, (a, b) in enumerate(actions, start=1):
        at.rows[i].cells[0].text = a
        at.rows[i].cells[1].text = b
    doc.add_page_break()
    chapter += 1

    h(f"{chapter}. Troubleshooting", 1, bookmark=f"sec_{chapter}")
    rows = [
        ("Cannot sign in", "Confirm email/password; use Forgot Password; verify role in hr_signup"),
        ("Redirected away from a page", "Role guards send users to their home panel (/dashboard, /head-hr, /ceo)"),
        ("Job not on /jobs", "Ensure the job is enabled/published"),
        ("CEO cannot create jobs", "Expected — CEO panel is read-only"),
        ("Staff cannot Apply", "Expected — signed-in staff are blocked from public apply"),
        ("Match score unexpected", "Confirm mandatory skills on JD and resume content"),
    ]
    tt = doc.add_table(rows=1 + len(rows), cols=2)
    tt.style = "Table Grid"
    tt.rows[0].cells[0].text = "Problem"
    tt.rows[0].cells[1].text = "What to try"
    for i, (a, b) in enumerate(rows, start=1):
        tt.rows[i].cells[0].text = a
        tt.rows[i].cells[1].text = b
    doc.add_page_break()
    chapter += 1

    h(f"{chapter}. FAQs", 1, bookmark=f"sec_{chapter}")
    faqs = [
        ("Is there a Super Admin role?", "No separate SUPER_ADMIN enum. Admin Login authenticates RECRUITER, HEAD_HR, or CEO."),
        ("Do candidates create accounts?", "Public apply is passwordless via /jobs. Candidate JWT sessions are cleared."),
        ("Can CEO edit jobs?", "No — CEO routes share Head HR pages in read-only mode."),
        ("Where is Feedback Admin?", "Recruiter navbar → Feedback (Admin) at /admin/feedback."),
        ("Where is HRMS Testing Feedback?", "Public Support menu → /support/hrms-feedback."),
    ]
    for q, a in faqs:
        p(f"Q: {q}", bold=True)
        p(f"A: {a}")
    doc.add_page_break()
    chapter += 1

    h(f"{chapter}. Appendix", 1, bookmark=f"sec_{chapter}")
    p("From repository root with the app running, regenerate this manual with:")
    bullet("python docs/user-manual/capture.py")
    bullet("python docs/user-manual/build.py")

    path = ROOT / f"{DOC_NAME}.docx"
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass
    doc.save(str(path))
    return path


def export_pdf_from_docx(docx_path: Path) -> Path:
    """Export PDF from the Word document so DOCX and PDF stay identical."""
    pdf_path = ROOT / f"{DOC_NAME}.pdf"
    try:
        import win32com.client  # type: ignore
    except ImportError as exc:
        raise SystemExit(
            "win32com/pywin32 is required to keep Word and PDF in sync. "
            "Install with: pip install pywin32"
        ) from exc

    word = None
    doc = None
    try:
        try:
            pdf_path.unlink(missing_ok=True)
        except OSError:
            pass
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.Fields.Update()
        if doc.TablesOfContents.Count >= 1:
            doc.TablesOfContents(1).Update()
        # 17 = wdFormatPDF
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path.resolve()),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            BitmapMissingFonts=True,
            DocStructureTags=True,
            CreateBookmarks=1,
        )
        doc.Save()
        print("PDF exported from Word (identical content).")
        return pdf_path
    except Exception as exc:  # noqa: BLE001
        raise SystemExit(f"Failed to export PDF from Word: {exc}") from exc
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def main():
    manifest = load_manifest()
    for legacy in ("User_Manual.docx", "User_Manual.pdf"):
        try:
            (ROOT / legacy).unlink(missing_ok=True)
        except OSError:
            pass
    docx = build_docx(manifest)
    print(f"DOCX: {docx}")
    pdf = export_pdf_from_docx(docx)
    print(f"PDF:  {pdf}")
    print(f"Figures: {len(manifest)}")
    print("Word and PDF are synced from the same DOCX source.")


if __name__ == "__main__":
    main()

"""
Text Extraction from PDF, DOCX, and image files.

Prefers PyMuPDF for digital PDF text; falls back to OCR for scanned/image PDFs
and direct image uploads. OCR uses RapidOCR (pip-only via requirements.txt);
optional system Tesseract is a secondary fallback. Optional PARSING_API is last resort.
"""
from __future__ import annotations

import io
import logging
import os
import re
import shutil
import threading
from typing import Any

import requests
from docx import Document

from app.core.timing import timing

logger = logging.getLogger(__name__)

PARSING_API_URL = os.getenv('PARSING_API_URL', 'http://localhost:4000')
PARSING_API_KEY = os.getenv('PARSING_API_KEY', 'your-api-key-here')
PARSING_API_FALLBACK = os.getenv('PARSING_API_FALLBACK', 'true').lower() in (
    '1',
    'true',
    'yes',
)
# Connect timeout kept short so a dead localhost:4000 fails fast in bulk.
PARSING_API_CONNECT_TIMEOUT = float(os.getenv('PARSING_API_CONNECT_TIMEOUT', '1.5'))
PARSING_API_READ_TIMEOUT = float(os.getenv('PARSING_API_READ_TIMEOUT', '30'))
PDF_MAX_PAGES = max(0, int(os.getenv('PDF_MAX_PAGES', '0')))
OCR_ENABLED = os.getenv('OCR_ENABLED', 'true').lower() in ('1', 'true', 'yes')
OCR_LANG = os.getenv('OCR_LANG', 'eng')
OCR_DPI = max(72, int(os.getenv('OCR_DPI', '250')))
# Try a lower DPI first for speed; escalate when OCR text is thin.
OCR_DPI_FAST = max(72, int(os.getenv('OCR_DPI_FAST', '180')))
MIN_TEXT_CHARS = 30
# Per-page digital text below this triggers OCR.
PAGE_OCR_TEXT_THRESHOLD = 80
# If digital text is below this and page has images, prefer OCR.
PAGE_SPARSE_TEXT_WITH_IMAGES = 200
RESUME_LAYOUT_ENABLED = os.getenv('RESUME_LAYOUT_ENABLED', 'true').lower() in (
    '1',
    'true',
    'yes',
)
JD_LAYOUT_ENABLED = os.getenv('JD_LAYOUT_ENABLED', 'true').lower() in (
    '1',
    'true',
    'yes',
)

IMAGE_EXTENSIONS = frozenset({'png', 'jpg', 'jpeg', 'webp', 'tif', 'tiff', 'bmp'})

# Invisible / formatting chars common in Word→PDF exports (ZWSP, soft hyphen, BOM, etc.)
_INVISIBLE_CHARS_RE = __import__('re').compile(
    r'[\u200b\u200c\u200d\u2060\ufeff\u00ad\u180e]'
)

_FIELD_LABEL_RE = __import__('re').compile(
    r'(?i)^(job\s*title|title|position|designation|role|location|work\s*location|'
    r'experience|work\s*experience|exp\.?|salary|ctc|compensation|employment\s*type|'
    r'job\s*type|company|department|skills?|required\s*skills?|primary\s*skills?|'
    r'qualification|notice\s*period|reports?\s*to)\b'
)


def _serialize_table_row(cells: list[str]) -> str:
    """Turn table cells into Label: Value lines for field extractors."""
    cleaned = [(__import__('re').sub(r'\s+', ' ', (c or '').strip())) for c in cells]
    cleaned = [c for c in cleaned if c]
    if not cleaned:
        return ''
    if len(cleaned) == 1:
        return cleaned[0]
    label, *rest = cleaned
    value = ' | '.join(rest)
    if _FIELD_LABEL_RE.match(label) or label.endswith(':') or len(label.split()) <= 4:
        label = label.rstrip(':').strip()
        return f'{label}: {value}'
    return ' | '.join(cleaned)


def _dedupe_append(parts: list[str], block: str) -> None:
    block = (block or '').strip()
    if not block:
        return
    norm = __import__('re').sub(r'\s+', ' ', block.lower())
    for existing in parts:
        if norm and norm in __import__('re').sub(r'\s+', ' ', existing.lower()):
            return
    parts.append(block)


def _extract_pdf_page_tables(page) -> str:
    """Serialize PyMuPDF find_tables() rows as Label: Value lines."""
    try:
        finder = page.find_tables()
    except Exception:
        return ''
    tables = getattr(finder, 'tables', None) or []
    lines: list[str] = []
    for table in tables:
        try:
            rows = table.extract() or []
        except Exception:
            continue
        for row in rows:
            if not row:
                continue
            cells = [str(c or '') for c in row]
            serialized = _serialize_table_row(cells)
            if serialized:
                lines.append(serialized)
    return '\n'.join(lines).strip()


def normalize_extracted_text(text: str) -> str:
    """
    Strip PDF/Word invisible characters that break name/header matching.
    Zero-width spaces after names (e.g. 'DHRUTI JADEJA\\u200b') are common.
    """
    if not text:
        return ''
    t = _INVISIBLE_CHARS_RE.sub('', text)
    t = t.replace('\xa0', ' ').replace('\u202f', ' ')
    return t


_rapidocr_engine: Any = None
_rapidocr_lock = threading.Lock()
# Cached probe: (available, engine_name_or_reason)
_ocr_engine_status: tuple[bool, str] | None = None


def _tesseract_available() -> bool:
    return bool(shutil.which('tesseract'))


def reset_ocr_engine_status_cache() -> None:
    """Test helper — clear the OCR availability probe cache."""
    global _ocr_engine_status
    _ocr_engine_status = None


def get_ocr_engine_status() -> tuple[bool, str]:
    """
    Return (available, detail). Probes once per process.

    available=True when RapidOCR import works or system Tesseract + pytesseract exist.
    """
    global _ocr_engine_status
    if _ocr_engine_status is not None:
        return _ocr_engine_status
    if not OCR_ENABLED:
        _ocr_engine_status = (False, 'OCR_ENABLED=false')
        return _ocr_engine_status
    try:
        import rapidocr_onnxruntime  # noqa: F401

        _ocr_engine_status = (True, 'rapidocr')
        return _ocr_engine_status
    except ImportError:
        pass
    except Exception as exc:
        logger.debug('RapidOCR probe failed: %s', exc)
    if _tesseract_available():
        try:
            import pytesseract  # noqa: F401

            _ocr_engine_status = (True, 'tesseract')
            return _ocr_engine_status
        except ImportError:
            _ocr_engine_status = (
                False,
                'tesseract binary found but pytesseract is not installed',
            )
            return _ocr_engine_status
    _ocr_engine_status = (
        False,
        'RapidOCR not installed and Tesseract unavailable '
        '(pip install rapidocr-onnxruntime, or install system Tesseract)',
    )
    return _ocr_engine_status


def ocr_engines_available() -> bool:
    """True when at least one local OCR engine can run."""
    ok, _ = get_ocr_engine_status()
    return ok


def ocr_unavailable_reason() -> str:
    _, detail = get_ocr_engine_status()
    return detail


def _get_rapidocr_engine() -> Any:
    """Lazy-load RapidOCR once (shared by extract + layout detections)."""
    global _rapidocr_engine
    if _rapidocr_engine is not None:
        return _rapidocr_engine
    from rapidocr_onnxruntime import RapidOCR

    with _rapidocr_lock:
        if _rapidocr_engine is None:
            _rapidocr_engine = RapidOCR()
        return _rapidocr_engine


def _ocr_with_rapidocr(image_bytes: bytes) -> str:
    """OCR via RapidOCR (pip-installable, no system binary)."""
    import numpy as np
    from PIL import Image

    from app.ai.parser.layout.preprocess import preprocess_image_bytes

    processed = preprocess_image_bytes(image_bytes)
    image = Image.open(io.BytesIO(processed))
    if image.mode not in ('RGB', 'L'):
        image = image.convert('RGB')
    arr = np.array(image)
    engine = _get_rapidocr_engine()
    result, _ = engine(arr)
    if not result:
        return ''
    return _join_ocr_detections_reading_order(result)


def _box_sort_key(box: Any) -> tuple[float, float]:
    """Sort key from RapidOCR box: top-to-bottom, then left-to-right."""
    try:
        # box is usually [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
        ys = [float(p[1]) for p in box]
        xs = [float(p[0]) for p in box]
        y = min(ys)
        x = min(xs)
        # Bucket Y so nearby same-line detections stay left-to-right
        return (round(y / 12.0) * 12.0, x)
    except Exception:
        return (0.0, 0.0)


def _join_ocr_detections_reading_order(detections: list) -> str:
    """Join RapidOCR [box, text, score] items in reading order."""
    rows: list[tuple[tuple[float, float], str]] = []
    for item in detections or []:
        if not isinstance(item, (list, tuple)) or len(item) < 2:
            continue
        text = item[1]
        if not text or not str(text).strip():
            continue
        box = item[0] if len(item) >= 1 else None
        rows.append((_box_sort_key(box), str(text).strip()))
    rows.sort(key=lambda r: r[0])
    return '\n'.join(text for _, text in rows).strip()


def _ocr_with_tesseract(image_bytes: bytes, *, lang: str | None = None) -> str:
    """OCR via system Tesseract + pytesseract (optional fallback)."""
    import pytesseract
    from PIL import Image

    from app.ai.parser.layout.preprocess import preprocess_image_bytes

    processed = preprocess_image_bytes(image_bytes)
    image = Image.open(io.BytesIO(processed))
    if image.mode not in ('RGB', 'L'):
        image = image.convert('RGB')
    text = pytesseract.image_to_string(image, lang=lang or OCR_LANG)
    return (text or '').strip()


def _ocr_image_bytes_plain(image_bytes: bytes, *, lang: str | None = None) -> str:
    """Run RapidOCR then Tesseract without layout structuring."""
    errors: list[str] = []

    try:
        text = _ocr_with_rapidocr(image_bytes)
        if text:
            return text
        errors.append('RapidOCR returned empty text')
    except ImportError as exc:
        errors.append(
            f'RapidOCR not installed ({exc}). Run: pip install -r requirements.txt'
        )
    except Exception as exc:
        errors.append(f'RapidOCR failed: {exc}')
        logger.warning('RapidOCR failed, trying Tesseract if available: %s', exc)

    if _tesseract_available():
        try:
            text = _ocr_with_tesseract(image_bytes, lang=lang)
            if text:
                return text
            errors.append('Tesseract returned empty text')
        except Exception as exc:
            errors.append(f'Tesseract failed: {exc}')

    detail = '; '.join(errors) if errors else 'no OCR engine available'
    raise ValueError(
        f'OCR failed ({detail}). Install RapidOCR with Python 3.10–3.12 '
        f'(pip install rapidocr-onnxruntime), or install system Tesseract as a fallback.'
    )


def _png_has_ink(image_bytes: bytes) -> bool:
    """True when a rendered page has enough contrast to be worth OCR.

    Blank / near-white pages still cost 10–15s of RapidOCR today. Detecting
    them after the cheap render (~100ms) does not skip Extract / Layout stages.
    """
    if not image_bytes or len(image_bytes) < 64:
        return False
    try:
        from PIL import Image

        image = Image.open(io.BytesIO(image_bytes))
        gray = image.convert('L')
        gray.thumbnail((240, 240))
        extrema = gray.getextrema()
        if extrema is None:
            return False
        lo, hi = int(extrema[0]), int(extrema[1])
        if hi - lo < 18:
            return False
        hist = gray.histogram()
        pixels = gray.size[0] * gray.size[1]
        if pixels <= 0:
            return False
        # Ink: darker than light-gray. Blank scans are almost all 240–255.
        dark = sum(hist[:200])
        return dark > pixels * 0.004
    except Exception:
        return True


def _ocr_image_bytes(image_bytes: bytes, *, lang: str | None = None) -> str:
    """
    Run OCR on raw image bytes.

    Primary: RapidOCR (+ OpenCV preprocess, optional layout).
    Secondary: system Tesseract if available.
    Fail-fast when no OCR engine is installed (avoids per-page layout retries).
    """
    if not OCR_ENABLED:
        raise ValueError('OCR is disabled (OCR_ENABLED=false)')

    if not ocr_engines_available():
        raise ValueError(
            f'OCR engines unavailable: {ocr_unavailable_reason()}. '
            'Install RapidOCR with Python 3.10–3.12 '
            '(pip install rapidocr-onnxruntime), or install system Tesseract.'
        )

    if RESUME_LAYOUT_ENABLED:
        try:
            from app.ai.parser.layout.detector import ocr_image_with_layout

            text, source = ocr_image_with_layout(
                image_bytes,
                ocr_fn=lambda b: _ocr_image_bytes_plain(b, lang=lang),
            )
            if text and text.strip():
                logger.debug('Layout OCR source=%s chars=%s', source, len(text))
                return text.strip()
            # Layout already ran RapidOCR. A second plain pass on the same
            # image is what made blank pages take ~12s each.
            raise ValueError(
                'OCR failed (RapidOCR returned empty text). Install RapidOCR with '
                'Python 3.10–3.12 (pip install rapidocr-onnxruntime), or install '
                'system Tesseract as a fallback.'
            )
        except ValueError:
            raise
        except Exception as exc:
            logger.warning('Layout OCR failed, falling back to plain OCR: %s', exc)

    return _ocr_image_bytes_plain(image_bytes, lang=lang)


def extract_text_from_image(file_data: bytes, filename: str = 'image.png') -> str:
    """Extract text from a standalone image via OCR."""
    text = _ocr_image_bytes(file_data)
    if len(text) < MIN_TEXT_CHARS:
        raise ValueError(
            f'OCR extracted only {len(text)} characters from image {filename}. '
            'The image may be blank, too low-resolution, or unreadable.'
        )
    logger.info('OCR extracted %s characters from image %s', len(text), filename)
    return text


def _render_page_png(page, dpi: int = OCR_DPI) -> bytes:
    """Render a PyMuPDF page to PNG bytes."""
    import fitz

    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    return pix.tobytes('png')


def _page_needs_ocr(page, digital_text: str) -> bool:
    """True when digital text is thin or the page is image-heavy with sparse text."""
    text_len = len((digital_text or '').strip())
    image_count = 0
    try:
        image_count = len(page.get_images(full=True) or [])
    except Exception:
        image_count = 0

    if text_len < PAGE_OCR_TEXT_THRESHOLD:
        return True
    if image_count > 0 and text_len < PAGE_SPARSE_TEXT_WITH_IMAGES:
        return True
    # Large image count with modest text often means scanned page + junk text layer
    if image_count >= 3 and text_len < 400:
        return True
    return False


def extract_text_from_pdf_pymupdf(file_data: bytes, *, dpi: int | None = None) -> str:
    """
    Extract text from PDF via PyMuPDF, with per-page OCR when needed.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ValueError('PyMuPDF (pymupdf) is not installed') from exc

    # Explicit dpi from caller wins; otherwise adaptive fast → full
    forced_dpi = max(72, int(dpi)) if dpi is not None else None
    fast_dpi = forced_dpi or max(72, min(OCR_DPI_FAST, OCR_DPI))
    full_dpi = forced_dpi or OCR_DPI

    doc = fitz.open(stream=file_data, filetype='pdf')
    try:
        page_count = len(doc)
        if PDF_MAX_PAGES:
            page_count = min(page_count, PDF_MAX_PAGES)

        text_parts: list[str] = []
        used_ocr = False
        max_dpi_used = fast_dpi
        ocr_ok = OCR_ENABLED and ocr_engines_available()
        ocr_skip_logged = False

        for page_num in range(page_count):
            page = doc[page_num]
            digital = (page.get_text('text') or '').strip()
            table_text = _extract_pdf_page_tables(page)

            if _page_needs_ocr(page, digital):
                if ocr_ok:
                    try:
                        png = _render_page_png(page, dpi=fast_dpi)
                        if not _png_has_ink(png):
                            logger.info(
                                'PDF page %s has no ink; skipping OCR (digital=%s)',
                                page_num + 1,
                                len(digital),
                            )
                            page_bits = []
                            if digital:
                                page_bits.append(digital)
                            if table_text:
                                _dedupe_append(page_bits, table_text)
                            if page_bits:
                                text_parts.append('\n\n'.join(page_bits))
                            continue
                        ocr_text = ''
                        try:
                            ocr_text = _ocr_image_bytes(png)
                        except ValueError as ocr_err:
                            logger.warning(
                                'OCR failed for PDF page %s: %s', page_num + 1, ocr_err
                            )
                        ocr_len = len((ocr_text or '').strip())
                        # Escalate DPI when the fast pass was thin, or empty on a
                        # page that still has ink (scanned page RapidOCR missed).
                        if (
                            forced_dpi is None
                            and full_dpi > fast_dpi
                            and ocr_len < PAGE_OCR_TEXT_THRESHOLD
                        ):
                            png = _render_page_png(page, dpi=full_dpi)
                            if _png_has_ink(png):
                                try:
                                    ocr_text = _ocr_image_bytes(png)
                                    max_dpi_used = max(max_dpi_used, full_dpi)
                                except ValueError as ocr_err:
                                    logger.warning(
                                        'OCR failed for PDF page %s at %sdpi: %s',
                                        page_num + 1,
                                        full_dpi,
                                        ocr_err,
                                    )
                        page_bits: list[str] = []
                        if ocr_text:
                            page_bits.append(ocr_text)
                            used_ocr = True
                        elif digital and len(digital) >= PAGE_OCR_TEXT_THRESHOLD:
                            page_bits.append(digital)
                        if table_text:
                            _dedupe_append(page_bits, table_text)
                        if page_bits:
                            text_parts.append('\n\n'.join(page_bits))
                        continue
                    except ValueError as ocr_err:
                        logger.warning(
                            'OCR failed for PDF page %s: %s', page_num + 1, ocr_err
                        )
                        # If engines disappeared mid-run, stop trying OCR on later pages
                        if not ocr_engines_available():
                            ocr_ok = False
                        page_bits = []
                        if digital and len(digital) >= PAGE_OCR_TEXT_THRESHOLD:
                            page_bits.append(digital)
                        if table_text:
                            _dedupe_append(page_bits, table_text)
                        if page_bits:
                            text_parts.append('\n\n'.join(page_bits))
                        continue
                else:
                    if OCR_ENABLED and not ocr_skip_logged:
                        logger.warning(
                            'OCR engines unavailable — skipping per-page OCR for this PDF (%s)',
                            ocr_unavailable_reason(),
                        )
                        ocr_skip_logged = True
                    page_bits = []
                    if digital and len(digital) >= PAGE_OCR_TEXT_THRESHOLD:
                        page_bits.append(digital)
                    elif digital:
                        # Keep whatever digital text exists; better than minutes of failed OCR
                        page_bits.append(digital)
                    if table_text:
                        _dedupe_append(page_bits, table_text)
                    if page_bits:
                        text_parts.append('\n\n'.join(page_bits))
            else:
                page_bits = []
                if digital:
                    page_bits.append(digital)
                if table_text:
                    _dedupe_append(page_bits, table_text)
                if page_bits:
                    text_parts.append('\n\n'.join(page_bits))

        extracted = '\n\n'.join(text_parts).strip()
        # Structure headers when resume or JD layout is enabled
        if (RESUME_LAYOUT_ENABLED or JD_LAYOUT_ENABLED) and extracted:
            try:
                from app.ai.parser.layout.detector import enhance_jd_text, enhance_resume_text, is_jd_layout_enabled

                if is_jd_layout_enabled():
                    structured = enhance_jd_text(extracted)
                else:
                    structured = enhance_resume_text(extracted)
                if structured and len(structured.strip()) >= MIN_TEXT_CHARS:
                    extracted = structured
            except Exception as exc:
                logger.debug('layout enhance skipped: %s', exc)

        if len(extracted) < MIN_TEXT_CHARS:
            raise ValueError(
                'Insufficient text extracted - PDF may be image-based or corrupted'
            )
        if used_ocr:
            logger.info(
                'Extracted %s characters from PDF (%s pages, OCR used, dpi=%s)',
                len(extracted),
                page_count,
                max_dpi_used,
            )
        return extracted
    finally:
        doc.close()


def extract_text_from_pdf_pypdf2(file_data: bytes) -> str:
    """Legacy PyPDF2 extraction (fallback when PyMuPDF unavailable)."""
    import PyPDF2

    pdf_file = io.BytesIO(file_data)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    pages = pdf_reader.pages
    if PDF_MAX_PAGES:
        pages = pages[:PDF_MAX_PAGES]
    text_parts = []
    for page_num, page in enumerate(pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        except Exception as e:
            logger.warning('Failed to extract text from page %s: %s', page_num + 1, e)
            continue
    extracted_text = '\n\n'.join(text_parts)
    if len(extracted_text.strip()) < MIN_TEXT_CHARS:
        raise ValueError('Insufficient text extracted - PDF may be image-based or corrupted')
    return extracted_text


def extract_text_from_pdf_via_api(file_data: bytes, filename: str) -> str:
    """Fallback: Extract text from PDF using external parsing API."""
    try:
        endpoint = f"{PARSING_API_URL}/api/v1/parse/resume"
        headers = {}
        if PARSING_API_KEY and PARSING_API_KEY != 'your-api-key-here':
            headers['X-API-Key'] = PARSING_API_KEY

        files = {'file': (filename, file_data, 'application/pdf')}
        response = requests.post(
            endpoint,
            files=files,
            headers=headers,
            timeout=(PARSING_API_CONNECT_TIMEOUT, PARSING_API_READ_TIMEOUT),
        )

        if response.status_code == 200:
            data = response.json()
            raw_text = data.get('raw_text', '')
            if raw_text and len(raw_text.strip()) >= MIN_TEXT_CHARS:
                logger.info(
                    'Successfully extracted %s characters via parsing API',
                    len(raw_text.strip()),
                )
                return raw_text
            raise ValueError('Parsing API returned insufficient text')
        error_data = (
            response.json()
            if response.headers.get('content-type', '').startswith('application/json')
            else {}
        )
        error_msg = error_data.get('error', f'Parsing API returned status {response.status_code}')
        raise ValueError(f'Parsing API error: {error_msg}')
    except requests.exceptions.RequestException as e:
        logger.warning('Parsing API request failed: %s', e)
        raise ValueError(f'Failed to extract text via API: {e}') from e
    except ValueError:
        raise
    except Exception as e:
        logger.warning('Parsing API extraction failed: %s', e)
        raise ValueError(f'Failed to extract text via API: {e}') from e


def extract_text_from_pdf(file_data: bytes, *, dpi: int | None = None) -> str:
    """Extract text from PDF: PyMuPDF (+OCR) → PyPDF2 → raise."""
    try:
        return extract_text_from_pdf_pymupdf(file_data, dpi=dpi)
    except ValueError as e:
        # If PyMuPDF missing, try PyPDF2; otherwise re-raise for OCR/API fallback chain.
        if 'pymupdf' in str(e).lower() or 'PyMuPDF' in str(e):
            logger.warning('PyMuPDF unavailable, falling back to PyPDF2: %s', e)
            return extract_text_from_pdf_pypdf2(file_data)
        raise
    except Exception as e:
        # Unexpected PyMuPDF errors — try PyPDF2 before failing.
        logger.warning('PyMuPDF extraction error, trying PyPDF2: %s', e)
        try:
            return extract_text_from_pdf_pypdf2(file_data)
        except Exception:
            raise ValueError(f'Failed to extract text from PDF: {e}') from e


def extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from DOCX (paragraphs + Label: Value tables + XML fallback)."""
    text_parts: list[str] = []
    try:
        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or '').strip() for cell in row.cells]
                # Deduplicate merged-cell repeats common in python-docx
                deduped: list[str] = []
                for c in cells:
                    if not c:
                        continue
                    if deduped and deduped[-1] == c:
                        continue
                    deduped.append(c)
                serialized = _serialize_table_row(deduped)
                if serialized:
                    text_parts.append(serialized)
    except Exception as e:
        logger.warning('python-docx extraction failed, trying XML fallback: %s', e)

    joined = '\n\n'.join(text_parts).strip()
    if len(joined) >= MIN_TEXT_CHARS:
        return joined

    # Fallback: unzip word/document.xml (handles some corrupted / odd DOCX)
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(io.BytesIO(file_data)) as zf:
            xml_name = 'word/document.xml'
            if xml_name not in zf.namelist():
                raise ValueError('DOCX missing word/document.xml')
            root = ET.fromstring(zf.read(xml_name))
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            chunks = []
            for node in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if node.text:
                    chunks.append(node.text)
                if node.tail:
                    chunks.append(node.tail)
            # Also collect paragraph breaks roughly
            xml_text = ' '.join(chunks)
            xml_text = re.sub(r'[ \t]+', ' ', xml_text)
            xml_text = re.sub(r'(\s*\n\s*)+', '\n', xml_text).strip()
            if len(xml_text) >= MIN_TEXT_CHARS:
                return xml_text
    except Exception as xml_err:
        logger.warning('DOCX XML fallback failed: %s', xml_err)

    if joined:
        return joined
    raise ValueError('Failed to extract text from DOCX: empty document')


def _force_pdf_ocr(file_data: bytes, *, dpi: int = 300) -> str:
    """Render every PDF page and OCR — last resort for scanned / empty digital layers."""
    if not ocr_engines_available():
        raise ValueError(
            f'Cannot force PDF OCR: {ocr_unavailable_reason()}'
        )
    try:
        import fitz
    except ImportError as exc:
        raise ValueError('PyMuPDF (pymupdf) is not installed') from exc
    doc = fitz.open(stream=file_data, filetype='pdf')
    try:
        page_count = len(doc)
        if PDF_MAX_PAGES:
            page_count = min(page_count, PDF_MAX_PAGES)
        parts: list[str] = []
        for page_num in range(page_count):
            page = doc[page_num]
            png = _render_page_png(page, dpi=dpi)
            ocr_text = _ocr_image_bytes(png)
            if ocr_text and ocr_text.strip():
                parts.append(ocr_text.strip())
        return '\n\n'.join(parts).strip()
    finally:
        doc.close()


@timing
def extract_text(file_data: bytes, filename: str, *, dpi: int | None = None) -> str:
    """
    Extract text from file based on extension.
    Tries local extraction (with OCR) first, falls back to parsing API for PDFs.
    Optional dpi overrides PDF OCR render resolution (e.g. bulk retry at 300).
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext in IMAGE_EXTENSIONS:
        text = extract_text_from_image(file_data, filename)
    elif ext == 'pdf':
        try:
            text = extract_text_from_pdf(file_data, dpi=dpi)
        except ValueError as e:
            error_msg = str(e)
            if 'Insufficient text' in error_msg or 'Failed to extract' in error_msg or 'OCR' in error_msg:
                if not PARSING_API_FALLBACK:
                    raise
                logger.info('Local PDF extraction failed, trying parsing API fallback...')
                try:
                    text = extract_text_from_pdf_via_api(file_data, filename)
                except Exception as api_error:
                    raise ValueError(
                        f'Local extraction failed: {error_msg}. '
                        f'API fallback also failed: {api_error}'
                    ) from api_error
            else:
                raise
        # Empty digital layer / failed OCR → force full-page OCR (only when engines exist)
        if len((text or '').strip()) < MIN_TEXT_CHARS:
            if ocr_engines_available():
                try:
                    forced = _force_pdf_ocr(file_data, dpi=max(dpi or 0, 300))
                    if len(forced) >= MIN_TEXT_CHARS:
                        text = forced
                except Exception as force_err:
                    logger.warning('Force PDF OCR failed: %s', force_err)
            else:
                logger.warning(
                    'Skipping force PDF OCR — %s',
                    ocr_unavailable_reason(),
                )
    elif ext == 'doc':
        raise ValueError('Legacy .doc format is not supported. Please use DOCX or PDF.')
    elif ext == 'docx':
        text = extract_text_from_docx(file_data)
    else:
        raise ValueError(f'Unsupported file type: {ext}')

    if RESUME_LAYOUT_ENABLED and text and ext != 'pdf':
        # PDF path already enhances inside pymupdf extractor
        try:
            from app.ai.parser.layout.detector import enhance_resume_text

            text = enhance_resume_text(text)
        except Exception as exc:
            logger.debug('enhance_resume_text skipped: %s', exc)
    return normalize_extracted_text(text or '')

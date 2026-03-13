"""
Text Extraction from PDF and DOCX files
"""
import io
import os
from typing import Optional
import PyPDF2
from docx import Document
import requests

# Parsing API configuration
PARSING_API_URL = os.getenv('PARSING_API_URL', 'http://localhost:4000')
PARSING_API_KEY = os.getenv('PARSING_API_KEY', 'your-api-key-here')
# Max PDF pages to extract (0 = all). Limits long PDFs for faster parsing; 25 is plenty for resumes.
PDF_MAX_PAGES = max(0, int(os.getenv('PDF_MAX_PAGES', '0')))


def extract_text_from_pdf_via_api(file_data: bytes, filename: str) -> str:
    """
    Fallback: Extract text from PDF using parsing API's parse endpoint
    We call the parse endpoint and extract just the raw_text from the response
    
    Args:
        file_data: PDF file bytes
        filename: Original filename
    
    Returns:
        Extracted text
    """
    try:
        endpoint = f"{PARSING_API_URL}/api/v1/parse/resume"
        headers = {}
        if PARSING_API_KEY and PARSING_API_KEY != 'your-api-key-here':
            headers['X-API-Key'] = PARSING_API_KEY
        
        files = {'file': (filename, file_data, 'application/pdf')}
        response = requests.post(endpoint, files=files, headers=headers, timeout=60)
        
        if response.status_code == 200:
            data = response.json()
            # Extract raw_text from the parsing response
            raw_text = data.get('raw_text', '')
            if raw_text and len(raw_text.strip()) >= 30:
                print(f"[INFO] Successfully extracted {len(raw_text.strip())} characters via parsing API")
                return raw_text
            else:
                raise ValueError("Parsing API returned insufficient text")
        else:
            error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
            error_msg = error_data.get('error', f"Parsing API returned status {response.status_code}")
            raise ValueError(f"Parsing API error: {error_msg}")
    except requests.exceptions.RequestException as e:
        print(f"[WARN] Parsing API request failed: {str(e)}")
        raise ValueError(f"Failed to extract text via API: {str(e)}")
    except Exception as e:
        print(f"[WARN] Parsing API extraction failed: {str(e)}")
        raise ValueError(f"Failed to extract text via API: {str(e)}")


def extract_text_from_pdf(file_data: bytes) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_data: PDF file bytes
    
    Returns:
        Extracted text
    """
    try:
        pdf_file = io.BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pages = pdf_reader.pages
        if PDF_MAX_PAGES:
            pages = pages[:PDF_MAX_PAGES]
        text_parts = []
        for page_num, page in enumerate(pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
            except Exception as e:
                print(f"[WARN] Failed to extract text from page {page_num + 1}: {str(e)}")
                continue
        extracted_text = '\n\n'.join(text_parts)
        
        # If we got very little text, it might be an image-based PDF
        # Lowered threshold to 30 characters to be more lenient
        if len(extracted_text.strip()) < 30:
            print(f"[WARN] Extracted only {len(extracted_text.strip())} characters, PDF might be image-based")
            raise ValueError("Insufficient text extracted - PDF may be image-based or corrupted")
        
        return extracted_text
    except ValueError:
        # Re-raise ValueError (our custom error)
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_data: bytes) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_data: DOCX file bytes
    
    Returns:
        Extracted text
    """
    try:
        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_data: bytes, filename: str) -> str:
    """
    Extract text from file based on extension
    Tries local extraction first, falls back to parsing API if needed
    
    Args:
        file_data: File bytes
        filename: Original filename
    
    Returns:
        Extracted text
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    
    if ext == 'pdf':
        try:
            # Try local extraction first
            return extract_text_from_pdf(file_data)
        except ValueError as e:
            error_msg = str(e)
            # If local extraction failed or returned insufficient text, try API fallback
            if 'Insufficient text' in error_msg or 'Failed to extract' in error_msg:
                print(f"[INFO] Local PDF extraction failed, trying parsing API fallback...")
                try:
                    return extract_text_from_pdf_via_api(file_data, filename)
                except Exception as api_error:
                    # If API also fails, raise the original error
                    raise ValueError(f"Local extraction failed: {error_msg}. API fallback also failed: {str(api_error)}")
            else:
                raise
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(file_data)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


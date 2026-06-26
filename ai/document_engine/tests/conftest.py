"""Shared test fixtures."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Senior Software Engineer with 8 years experience.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    table.cell(1, 0).text = "Skill"
    table.cell(1, 1).text = "SQL"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    path = tmp_path / "sample.txt"
    path.write_text("Plain text resume content for testing extraction.\n", encoding="utf-8")
    return path
